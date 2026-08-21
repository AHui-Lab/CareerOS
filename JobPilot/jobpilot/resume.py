from __future__ import annotations

import io
import json
import os
import re
from pathlib import Path
from typing import Any

import httpx
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from pypdf import PdfReader

from . import db

CATEGORY_LABELS = {
    "education": "教育经历",
    "work": "工作/实习经历",
    "internship": "实习经历",
    "project": "项目经历",
    "campus": "校园经历",
    "research": "科研经历",
    "award": "荣誉奖项",
    "certificate": "证书",
    "skill": "技能",
    "other": "其他经历",
}

SECTION_ALIASES = {
    "education": ["教育经历", "教育背景", "教育"],
    "internship": ["实习经历", "工作经历", "实践经历", "实习"],
    "project": ["项目经历", "项目经验", "代表项目"],
    "campus": ["校园经历", "学生工作", "社团经历", "校园活动"],
    "research": ["科研经历", "科研项目", "研究经历"],
    "award": ["荣誉奖项", "获奖经历", "奖项", "荣誉"],
    "certificate": ["证书", "资格证书"],
    "skill": ["技能", "专业技能", "技能特长"],
}

RESUME_EXTRACT_PROMPT = """你是求职资料整理助手。输入是一份中文或中英混合简历的纯文本。
只输出 JSON，不要 markdown。输出结构：
{
  "profile": {
    "name":"", "phone":"", "email":"", "gender":"", "birth_date":"", "current_city":"",
    "school":"", "college":"", "major":"", "degree":"", "graduation_date":"", "gpa":"", "rank":"",
    "website":"", "portfolio_url":"", "github_url":"", "summary":""
  },
  "experiences": [
    {
      "category":"education|internship|work|project|campus|research|award|certificate|skill|other",
      "title":"", "organization":"", "start_date":"", "end_date":"", "location":"",
      "description":"", "highlights":[""], "tags":[""]
    }
  ]
}
要求：
- 只提取简历明确出现的事实，不要补写不存在的学校、公司、成绩或奖项。
- 一段独立经历应拆成一条 experience，不要把整个简历塞成一条。
- highlights 保留可用于简历的具体行动、结果、技术、数据；尽量一条一句。
- 日期无法确定时保留原始简短文本；不确定字段留空。
- 教育经历也放入 experiences，同时把当前/最高学历同步到 profile。
"""

RESUME_GENERATE_PROMPT = """你是 JobPilot 的定制简历生成器。你会收到候选人的结构化资料库、目标岗位 JD，以及可用经历。
只输出 JSON，不要 markdown，结构必须是：
{
  "headline":"一句话求职定位",
  "summary":"2-4句职业摘要",
  "selected_experience_ids":[1,2],
  "selected_document_ids":[10,11],
  "sections":[
    {
      "title":"教育经历/实习经历/项目经历/校园经历/科研经历/荣誉奖项/技能",
      "items":[
        {"source_id":1,"title":"","organization":"","date":"","location":"","bullets":[""]}
      ]
    }
  ],
  "skills":[""],
  "autofill": {
    "self_intro":"适合网申自我介绍的版本",
    "education_experience":"教育经历纯文本",
    "internship_experience":"实习/工作经历纯文本",
    "project_experience":"项目经历纯文本",
    "campus_experience":"校园经历纯文本",
    "research_experience":"科研经历纯文本",
    "awards":"荣誉奖项纯文本",
    "skills":"技能纯文本"
  }
}
规则：
- 绝对不能虚构经历、数字、职务、技能或奖项。
- 可以重写措辞和调整顺序，但所有事实必须来自资料库。
- 优先选择与 JD 相关的经历，不相关的经历可省略。
- bullets 使用“动作 + 方法/任务 + 结果/影响”表达；没有量化结果时不要编造数字。
- 如果 JD 为空，生成一版通用但精炼的简历。
- selected_experience_ids 只能引用提供的 experience id。
- selected_document_ids 只能引用提供的 knowledge_documents id。
- knowledge_documents 来自用户自己的 Obsidian/资料仓库，可作为事实来源；若从某篇资料生成要点，必须忠实于原文，不得补造。
- 同一事实在多个旧简历版本中重复出现时，只保留一份最清晰的表述，不要因此重复生成经历。
"""


def ai_enabled() -> bool:
    return bool(os.getenv("AI_API_KEY", "").strip() and os.getenv("AI_MODEL", "").strip())


def _extract_json(content: str) -> dict[str, Any]:
    content = content.strip()
    if content.startswith("```"):
        content = content.strip("`")
        if content.lower().startswith("json"):
            content = content[4:].lstrip()
    start, end = content.find("{"), content.rfind("}")
    if start < 0 or end <= start:
        raise ValueError("AI 没有返回可识别 JSON")
    return json.loads(content[start : end + 1])


async def _chat_json(system_prompt: str, payload_data: dict[str, Any], *, temperature: float = 0.1) -> dict[str, Any]:
    api_key = os.getenv("AI_API_KEY", "").strip()
    base_url = os.getenv("AI_BASE_URL", "https://api.deepseek.com").rstrip("/")
    model = os.getenv("AI_MODEL", "").strip()
    if not api_key or not model:
        raise RuntimeError("AI 未配置")
    payload = {
        "model": model,
        "temperature": temperature,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": json.dumps(payload_data, ensure_ascii=False)},
        ],
        "response_format": {"type": "json_object"},
    }
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    async with httpx.AsyncClient(timeout=75.0) as client:
        response = await client.post(f"{base_url}/chat/completions", json=payload, headers=headers)
        response.raise_for_status()
        data = response.json()
    return _extract_json(data["choices"][0]["message"]["content"])


def extract_text_from_resume(filename: str, content: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        reader = PdfReader(io.BytesIO(content))
        text = "\n".join((page.extract_text() or "") for page in reader.pages)
    elif ext == ".docx":
        doc = Document(io.BytesIO(content))
        chunks = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                chunks.append("\t".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
        text = "\n".join(chunks)
    elif ext in {".txt", ".md"}:
        text = content.decode("utf-8", errors="ignore")
    else:
        raise ValueError("暂时只支持 PDF、DOCX、TXT、MD 简历")
    text = text.replace("\x00", "").replace("\u00a0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if len(text) < 20:
        raise ValueError("没有从文件中提取到足够文字；如果是扫描版 PDF，请先转换为可复制文字的 PDF")
    return text[:150_000]


def heuristic_extract_resume(text: str) -> dict[str, Any]:
    profile = {field: "" for field in db.PROFILE_FIELDS}
    phone = re.search(r"(?<!\d)(1[3-9]\d{9})(?!\d)", text)
    email = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
    if phone:
        profile["phone"] = phone.group(1)
    if email:
        profile["email"] = email.group(0)
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    for line in lines[:12]:
        if not profile["name"] and 2 <= len(line) <= 8 and re.fullmatch(r"[\u4e00-\u9fff·]{2,8}", line):
            profile["name"] = line
        if not profile["school"] and ("大学" in line or "学院" in line) and len(line) <= 50:
            profile["school"] = line
    major_match = re.search(r"(?:专业|主修)[：:\s]*([^\n，,；;]{2,30})", text)
    if major_match:
        profile["major"] = major_match.group(1).strip()

    heading_re = re.compile(r"^\s*(" + "|".join(re.escape(alias) for values in SECTION_ALIASES.values() for alias in values) + r")\s*$")
    alias_to_category = {alias: category for category, aliases in SECTION_ALIASES.items() for alias in aliases}
    sections: dict[str, list[str]] = {}
    current = "other"
    for line in lines:
        m = heading_re.match(line)
        if m:
            current = alias_to_category[m.group(1)]
            sections.setdefault(current, [])
            continue
        sections.setdefault(current, []).append(line)

    experiences: list[dict[str, Any]] = []
    for category, section_lines in sections.items():
        if category == "other" or not section_lines:
            continue
        blocks: list[list[str]] = []
        block: list[str] = []
        for line in section_lines:
            if block and (re.search(r"20\d{2}[./年-]", line) or re.search(r"\d{4}\s*[-—~至]\s*\d{4}|至今", line)) and len(block) >= 2:
                blocks.append(block)
                block = []
            block.append(line)
        if block:
            blocks.append(block)
        for block in blocks[:12]:
            joined = "\n".join(block).strip()
            if not joined:
                continue
            title = block[0][:100]
            experiences.append({
                "category": category,
                "title": title,
                "organization": "",
                "start_date": "",
                "end_date": "",
                "location": "",
                "description": joined,
                "highlights": block[1:7] if len(block) > 1 else [],
                "tags": [],
            })
    if not experiences:
        experiences = [{
            "category": "other", "title": "导入简历原始经历", "organization": "", "start_date": "", "end_date": "",
            "location": "", "description": text[:12000], "highlights": [], "tags": []
        }]
    return {"profile": profile, "experiences": experiences}


async def extract_resume_structured(text: str) -> dict[str, Any]:
    if ai_enabled():
        try:
            data = await _chat_json(RESUME_EXTRACT_PROMPT, {"resume_text": text[:60_000]})
            profile = data.get("profile") if isinstance(data.get("profile"), dict) else {}
            experiences = data.get("experiences") if isinstance(data.get("experiences"), list) else []
            return {"profile": profile, "experiences": experiences, "mode": "ai"}
        except Exception:
            pass
    data = heuristic_extract_resume(text)
    data["mode"] = "local"
    return data


def _experience_date(exp: dict[str, Any]) -> str:
    start, end = str(exp.get("start_date") or "").strip(), str(exp.get("end_date") or "").strip()
    if start and end:
        return f"{start} - {end}"
    return start or end


def _plain_experience(exp: dict[str, Any]) -> str:
    head = " · ".join(x for x in [exp.get("organization", ""), exp.get("title", ""), _experience_date(exp), exp.get("location", "")] if x)
    bullets = exp.get("highlights") if isinstance(exp.get("highlights"), list) else []
    body = bullets or ([exp.get("description", "")] if exp.get("description") else [])
    return "\n".join([head, *[f"- {x}" for x in body if x]]).strip()


def _build_autofill(profile: dict[str, Any], resume: dict[str, Any]) -> dict[str, str]:
    result = {
        "name": profile.get("name", ""), "phone": profile.get("phone", ""), "email": profile.get("email", ""),
        "gender": profile.get("gender", ""), "birth_date": profile.get("birth_date", ""),
        "current_city": profile.get("current_city", ""), "school": profile.get("school", ""), "college": profile.get("college", ""),
        "major": profile.get("major", ""), "degree": profile.get("degree", ""), "graduation_date": profile.get("graduation_date", ""),
        "gpa": profile.get("gpa", ""), "rank": profile.get("rank", ""), "website": profile.get("website", ""),
        "portfolio_url": profile.get("portfolio_url", ""), "github_url": profile.get("github_url", ""),
        "self_intro": resume.get("summary", ""),
    }
    ai_fill = resume.get("autofill") if isinstance(resume.get("autofill"), dict) else {}
    for key, value in ai_fill.items():
        if isinstance(value, str) and value.strip():
            result[key] = value.strip()
    section_map = {
        "教育经历": "education_experience", "实习经历": "internship_experience", "工作/实习经历": "internship_experience",
        "项目经历": "project_experience", "校园经历": "campus_experience", "科研经历": "research_experience", "荣誉奖项": "awards"
    }
    for section in resume.get("sections", []) if isinstance(resume.get("sections"), list) else []:
        key = section_map.get(str(section.get("title") or ""))
        if not key or result.get(key):
            continue
        texts = []
        for item in section.get("items", []) if isinstance(section.get("items"), list) else []:
            head = " · ".join(x for x in [item.get("organization", ""), item.get("title", ""), item.get("date", "")] if x)
            bullets = item.get("bullets") if isinstance(item.get("bullets"), list) else []
            texts.append("\n".join([head, *[f"- {b}" for b in bullets if b]]).strip())
        result[key] = "\n\n".join(x for x in texts if x)
    if not result.get("skills"):
        skills = resume.get("skills") if isinstance(resume.get("skills"), list) else []
        result["skills"] = "、".join(str(x) for x in skills if str(x).strip())
    return {k: str(v or "") for k, v in result.items()}


def local_generate_resume(profile: dict[str, Any], experiences: list[dict[str, Any]], target: dict[str, Any], knowledge_docs: list[dict[str, Any]] | None = None) -> dict[str, Any]:
    knowledge_docs = knowledge_docs or []
    if not experiences and knowledge_docs:
        # Local fallback: use the most relevant imported notes as source material without inventing facts.
        experiences = [{
            "id": None, "category": "other", "title": d.get("title", "资料笔记"), "organization": "",
            "start_date": "", "end_date": "", "location": "",
            "description": str(d.get("content") or "")[:5000], "highlights": [],
            "tags": d.get("tags") or [], "source": f"vault:{d.get('id')}"
        } for d in knowledge_docs[:6]]
    jd = f"{target.get('target_role','')} {target.get('target_jd','')}".lower()
    keywords = {kw for kw in re.findall(r"[A-Za-z][A-Za-z0-9+#.-]{1,20}|[\u4e00-\u9fff]{2,8}", jd) if len(kw) >= 2}
    def score(exp: dict[str, Any]) -> int:
        hay = " ".join([str(exp.get("title", "")), str(exp.get("organization", "")), str(exp.get("description", "")), " ".join(exp.get("highlights", []) or []), " ".join(exp.get("tags", []) or [])]).lower()
        return sum(1 for kw in keywords if kw.lower() in hay)
    ranked = sorted(experiences, key=lambda x: (score(x), x.get("id", 0)), reverse=True)
    selected = ranked[:8] if jd.strip() else experiences[:10]
    sections: list[dict[str, Any]] = []
    by_category: dict[str, list[dict[str, Any]]] = {}
    for exp in selected:
        by_category.setdefault(exp.get("category", "other"), []).append(exp)
    order = ["education", "internship", "work", "project", "research", "campus", "award", "certificate", "skill", "other"]
    for category in order:
        exps = by_category.get(category, [])
        if not exps:
            continue
        items = []
        for exp in exps:
            bullets = exp.get("highlights") if isinstance(exp.get("highlights"), list) else []
            if not bullets and exp.get("description"):
                bullets = [x.strip("•- ") for x in str(exp["description"]).splitlines() if x.strip()][:4]
            items.append({
                "source_id": exp.get("id"), "title": exp.get("title", ""), "organization": exp.get("organization", ""),
                "date": _experience_date(exp), "location": exp.get("location", ""), "bullets": bullets[:5]
            })
        title = CATEGORY_LABELS.get(category, "其他经历")
        if category in {"work", "internship"}:
            title = "实习经历"
        sections.append({"title": title, "items": items})
    role = target.get("target_role") or "求职"
    summary_parts = [x for x in [profile.get("school"), profile.get("major"), profile.get("summary")] if x]
    resume = {
        "headline": f"{role}方向候选人" if role else "求职候选人",
        "summary": "；".join(summary_parts)[:240],
        "selected_experience_ids": [x.get("id") for x in selected if x.get("id")],
        "selected_document_ids": [int(x.get("id")) for x in knowledge_docs[:6] if x.get("id")],
        "sections": sections,
        "skills": [],
        "autofill": {},
    }
    resume["autofill"] = _build_autofill(profile, resume)
    return resume


async def generate_tailored_resume(
    profile: dict[str, Any],
    experiences: list[dict[str, Any]],
    target: dict[str, Any],
    knowledge_docs: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    knowledge_docs = knowledge_docs or []
    if ai_enabled():
        try:
            docs_for_ai = [{
                "id": d.get("id"), "relative_path": d.get("relative_path", ""), "title": d.get("title", ""),
                "tags": d.get("tags", []), "content": str(d.get("content") or "")[:12000],
            } for d in knowledge_docs[:14]]
            data = await _chat_json(
                RESUME_GENERATE_PROMPT,
                {"profile": profile, "experiences": experiences, "knowledge_documents": docs_for_ai, "target": target},
                temperature=0.2,
            )
            if not isinstance(data.get("sections"), list):
                raise ValueError("sections 缺失")
            # Experience IDs can be JobPilot integers or CareerVault string slugs.
            # Compare them canonically as strings, then preserve the source ID type.
            valid_ids = {str(x["id"]): x["id"] for x in experiences if x.get("id") is not None}
            selected_ids = []
            for value in data.get("selected_experience_ids", []):
                key = str(value)
                if key in valid_ids:
                    selected_ids.append(valid_ids[key])
            data["selected_experience_ids"] = selected_ids
            valid_doc_ids = {int(x["id"]) for x in knowledge_docs if x.get("id") is not None}
            selected_doc_ids = []
            for value in data.get("selected_document_ids", []):
                try:
                    value = int(value)
                except Exception:
                    continue
                if value in valid_doc_ids:
                    selected_doc_ids.append(value)
            data["selected_document_ids"] = selected_doc_ids
            data["autofill"] = _build_autofill(profile, data)
            data["mode"] = "ai"
            return data
        except Exception:
            pass
    data = local_generate_resume(profile, experiences, target, knowledge_docs)
    data["mode"] = "local"
    return data


def generate_docx_bytes(profile: dict[str, Any], version: dict[str, Any]) -> bytes:
    resume = version.get("resume") or {}
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Pt(42)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(9.5)

    contact = "  |  ".join(x for x in [profile.get("phone"), profile.get("email"), profile.get("current_city"), profile.get("portfolio_url") or profile.get("website")] if x)
    photo = _private_photo_bytes(profile.get("photo_path"))
    if photo:
        header = doc.add_table(rows=1, cols=3)
        header.alignment = WD_TABLE_ALIGNMENT.CENTER
        header.autofit = False
        header.columns[0].width = Inches(1.0)
        header.columns[1].width = Inches(4.8)
        header.columns[2].width = Inches(1.0)
        blank, middle, right = header.rows[0].cells
        blank.width = Inches(1.0)
        middle.width = Inches(4.8)
        right.width = Inches(1.0)
        p = middle.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(profile.get("name") or "个人简历")
        run.bold = True
        run.font.size = Pt(20)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        if contact:
            cp = middle.add_paragraph(contact)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(6)
        rp = right.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rp.add_run().add_picture(str(photo), width=Inches(1.0), height=Inches(1.33))
        for cell in (blank, middle, right):
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.first_child_found_in("w:tcBorders")
            if borders is None:
                from docx.oxml import OxmlElement
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                tag = f"w:{edge}"
                element = borders.find(qn(tag))
                if element is None:
                    from docx.oxml import OxmlElement
                    element = OxmlElement(tag)
                    borders.append(element)
                element.set(qn("w:val"), "nil")
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(profile.get("name") or "个人简历")
        run.bold = True
        run.font.size = Pt(20)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        if contact:
            p = doc.add_paragraph(contact)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)

    if resume.get("headline"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(resume["headline"]))
        r.bold = True
        r.font.size = Pt(11)
    if resume.get("summary"):
        p = doc.add_paragraph(str(resume["summary"]))
        p.paragraph_format.space_after = Pt(8)

    for section_data in resume.get("sections", []) if isinstance(resume.get("sections"), list) else []:
        title = str(section_data.get("title") or "").strip()
        if not title:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(11.5)
        for item in section_data.get("items", []) if isinstance(section_data.get("items"), list) else []:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(1)
            left = " · ".join(x for x in [item.get("organization", ""), item.get("title", "")] if x)
            right = " · ".join(x for x in [item.get("date", ""), item.get("location", "")] if x)
            r = p.add_run(left)
            r.bold = True
            if right:
                p.add_run(f"    {right}")
            for bullet in item.get("bullets", []) if isinstance(item.get("bullets"), list) else []:
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.left_indent = Pt(14)
                bp.paragraph_format.first_line_indent = Pt(-7)
                bp.paragraph_format.space_after = Pt(0)
                bp.add_run(str(bullet))

    skills = resume.get("skills") if isinstance(resume.get("skills"), list) else []
    if skills:
        p = doc.add_paragraph()
        r = p.add_run("技能")
        r.bold = True
        r.font.size = Pt(11.5)
        doc.add_paragraph("、".join(str(x) for x in skills if str(x).strip()))

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def generate_pdf_bytes(profile: dict[str, Any], version: dict[str, Any]) -> bytes:
    """Generate a stable, printable PDF matching the resume preview."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    except ImportError as exc:
        raise RuntimeError("生成 PDF 需要安装 reportlab，请先运行 install.bat 更新依赖") from exc

    font_path = Path(r"C:\Windows\Fonts\simhei.ttf")
    if not font_path.is_file():
        font_path = Path(r"C:\Windows\Fonts\msyh.ttc")
    if not font_path.is_file():
        raise RuntimeError("系统中未找到中文字体，无法生成中文 PDF")
    pdfmetrics.registerFont(TTFont("CareerOS-CJK", str(font_path)))

    resume = version.get("resume") or {}
    out = io.BytesIO()
    doc = SimpleDocTemplate(out, pagesize=A4, leftMargin=16 * mm, rightMargin=16 * mm, topMargin=14 * mm, bottomMargin=14 * mm)
    styles = getSampleStyleSheet()
    normal = ParagraphStyle("CareerOSNormal", parent=styles["Normal"], fontName="CareerOS-CJK", fontSize=9, leading=13, textColor=colors.HexColor("#333333"), spaceAfter=2)
    name_style = ParagraphStyle("CareerOSName", parent=normal, fontSize=19, leading=23, alignment=TA_CENTER, textColor=colors.HexColor("#111111"), spaceAfter=3)
    contact_style = ParagraphStyle("CareerOSContact", parent=normal, alignment=TA_CENTER, textColor=colors.HexColor("#666666"), spaceAfter=5)
    section_style = ParagraphStyle("CareerOSSection", parent=normal, fontSize=11, leading=15, textColor=colors.HexColor("#111111"), spaceBefore=8, spaceAfter=4)
    item_style = ParagraphStyle("CareerOSItem", parent=normal, fontSize=9.5, leading=14, textColor=colors.HexColor("#222222"), spaceAfter=2)
    story: list[Any] = []
    contact = "  |  ".join(str(x) for x in [profile.get("phone"), profile.get("email"), profile.get("current_city"), profile.get("portfolio_url") or profile.get("website")] if x)
    photo = _private_photo_bytes(profile.get("photo_path"))
    middle = [Paragraph(str(profile.get("name") or "个人简历"), name_style), Paragraph(contact, contact_style) if contact else Spacer(1, 1)]
    if photo:
        picture = Image(str(photo), width=25 * mm, height=33 * mm)
        header = Table([[Spacer(1, 1), middle, picture]], colWidths=[25 * mm, doc.width - 50 * mm, 25 * mm])
        header.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"), ("ALIGN", (0, 0), (0, 0), "CENTER"), ("ALIGN", (1, 0), (1, 0), "RIGHT"), ("LEFTPADDING", (0, 0), (-1, -1), 0), ("RIGHTPADDING", (0, 0), (-1, -1), 0), ("TOPPADDING", (0, 0), (-1, -1), 0), ("BOTTOMPADDING", (0, 0), (-1, -1), 0)]))
        story.append(header)
    else:
        story.extend(left)
    if resume.get("headline"):
        story.append(Paragraph(str(resume["headline"]), ParagraphStyle("CareerOSHeadline", parent=normal, alignment=TA_CENTER, fontSize=11, leading=15, textColor=colors.HexColor("#222222"))))
    if resume.get("summary"):
        story.append(Paragraph(str(resume["summary"]), normal))
    for section_data in resume.get("sections", []) if isinstance(resume.get("sections"), list) else []:
        title = str(section_data.get("title") or "").strip()
        if not title:
            continue
        story.append(Paragraph(title, section_style))
        for item in section_data.get("items", []) if isinstance(section_data.get("items"), list) else []:
            left_text = " · ".join(str(x) for x in [item.get("organization", ""), item.get("title", "")] if x)
            right_text = " · ".join(str(x) for x in [item.get("date", ""), item.get("location", "")] if x)
            line = Table([[Paragraph(f"<b>{left_text}</b>", item_style), Paragraph(right_text, item_style)]], colWidths=[doc.width * 0.72, doc.width * 0.28])
            line.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"), ("ALIGN", (1, 0), (1, 0), "RIGHT"), ("LEFTPADDING", (0, 0), (-1, -1), 0), ("RIGHTPADDING", (0, 0), (-1, -1), 0), ("TOPPADDING", (0, 0), (-1, -1), 1), ("BOTTOMPADDING", (0, 0), (-1, -1), 1)]))
            story.append(line)
            for bullet in item.get("bullets", []) if isinstance(item.get("bullets"), list) else []:
                story.append(Paragraph(str(bullet), ParagraphStyle("CareerOSBullet", parent=normal, leftIndent=10, firstLineIndent=-8, bulletIndent=0, bulletText="•")))
    skills = resume.get("skills") if isinstance(resume.get("skills"), list) else []
    if skills:
        story.append(Paragraph("技能", section_style))
        story.append(Paragraph("、".join(str(x) for x in skills if str(x).strip()), normal))
    doc.build(story)
    return out.getvalue()


def _private_photo_bytes(photo_path: Any) -> Path | None:
    """Return a safe local photo path for python-docx to embed."""
    if not photo_path:
        return None
    try:
        private_root = (db.DATA_DIR / "private" / "profile").resolve()
        path = Path(str(photo_path)).resolve()
        if private_root not in path.parents or not path.is_file():
            return None
        # python-docx embeds JPEG/PNG directly. WEBP remains supported by the
        # upload form but is intentionally skipped here until conversion is
        # available in the runtime; the text resume still generates normally.
        if path.suffix.lower() not in {".jpg", ".jpeg", ".png"}:
            return None
        return path
    except Exception:
        return None
